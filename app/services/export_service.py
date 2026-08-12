"""
ChemAI 导出服务
试卷 Word 导出 + 报告 HTML 生成
"""
import io
from typing import List, Dict, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def export_exam_to_docx(
    questions: List[Dict],
    exam_name: str = "化学试卷",
    school_name: str = "",
    with_answers: bool = False,
) -> io.BytesIO:
    """
    生成 Word 格式试卷

    Args:
        questions: [{"content": "...", "options": [...], "answer": "...", "type": "choice/fill/...", "knowledge_points": [...]}]
        exam_name: 考试名称
        school_name: 学校名称
        with_answers: 是否包含答案

    Returns:
        BytesIO buffer with .docx content
    """
    doc = Document()

    # A4 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "SimSun"

    # 密封线（左侧）
    # 简化处理：直接写标题区

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{school_name}\n{exam_name}" if school_name else exam_name)
    run.bold = True
    run.font.size = Pt(16)

    # 信息区
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info.add_run(f"姓名：__________  班级：__________  得分：__________").font.size = Pt(11)

    doc.add_paragraph()  # 空行

    # 分类渲染题目
    qt_labels = {
        "choice": "一、选择题",
        "fill": "二、填空题",
        "calc": "三、计算题",
        "experiment": "四、实验题",
        "inference": "五、推断题",
    }

    # 按题型分组
    by_type: Dict[str, List[Dict]] = {}
    for q in questions:
        qt = q.get("type", "choice")
        by_type.setdefault(qt, []).append(q)

    q_num = 0
    for qt, label in qt_labels.items():
        qs = by_type.get(qt, [])
        if not qs:
            continue

        # 题型标题
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        run = h.add_run(label)
        run.bold = True
        run.font.size = Pt(13)

        for q in qs:
            q_num += 1
            content = q.get("content", "")

            if qt == "choice":
                # 选择题: 题号 + 题干 + 选项
                p = doc.add_paragraph()
                p.add_run(f"{q_num}. ").bold = True
                p.add_run(content)
                options = q.get("options", [])
                for opt in options:
                    doc.add_paragraph(f"    {opt}")

                if with_answers:
                    ans_p = doc.add_paragraph()
                    ans_run = ans_p.add_run(f"    答案: {q.get('answer', '')}")
                    ans_run.font.color.rgb = RGBColor(255, 0, 0)
                    ans_run.font.size = Pt(9)

            elif qt == "fill":
                p = doc.add_paragraph()
                p.add_run(f"{q_num}. ").bold = True
                p.add_run(content)

                if with_answers:
                    ans_p = doc.add_paragraph()
                    ans_run = ans_p.add_run(f"    答案: {q.get('answer', '')}")
                    ans_run.font.color.rgb = RGBColor(255, 0, 0)
                    ans_run.font.size = Pt(9)

            elif qt in ("calc", "experiment", "inference"):
                p = doc.add_paragraph()
                p.add_run(f"{q_num}. ").bold = True
                p.add_run(content)
                # 留答题空间
                doc.add_paragraph("    ")
                doc.add_paragraph("    ")

                if with_answers:
                    ans = q.get("answer", "")
                    analysis = q.get("analysis", "")
                    ans_p = doc.add_paragraph()
                    ans_run = ans_p.add_run(f"    答案: {ans}")
                    ans_run.font.color.rgb = RGBColor(255, 0, 0)
                    ans_run.font.size = Pt(9)
                    if analysis:
                        ans_p2 = doc.add_paragraph()
                        ans_r2 = ans_p2.add_run(f"    解析: {analysis[:200]}")
                        ans_r2.font.color.rgb = RGBColor(0, 100, 0)
                        ans_r2.font.size = Pt(9)

            # 题目间间距
            doc.add_paragraph()

    # 如果含答案，标记
    if with_answers:
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("（含答案版）")
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_report_html(
    exam_data: Dict,
    report_type: str = "teacher",  # "teacher" or "student"
) -> str:
    """
    生成报告 HTML

    Args:
        exam_data: 考试统计数据
        report_type: teacher 或 student
    """
    is_teacher = (report_type == "teacher")

    rows_html = ""
    top_errors = exam_data.get("question_stats", [])[:5]

    for q in top_errors:
        kps = ", ".join(q.get("knowledge_points", []))
        rows_html += f"""
        <tr>
            <td>{q.get("question_number", "?")}</td>
            <td>{kps}</td>
            <td>{q.get("error_count", 0)} 人</td>
            <td>{int(q.get("error_rate", 0) * 100)}%</td>
        </tr>"""

    teacher_section = ""
    if is_teacher:
        kp_rows = ""
        for kp in exam_data.get("knowledge_point_stats", [])[:5]:
            kp_rows += f"""
            <tr>
                <td>{kp.get("knowledge_point", "?")}</td>
                <td>{kp.get("error_count", 0)}</td>
                <td>{int(kp.get("error_rate", 0) * 100)}%</td>
            </tr>"""

        teacher_section = f"""
        <h2>知识点错误分布</h2>
        <table>
            <thead><tr><th>知识点</th><th>错误人数</th><th>错误率</th></tr></thead>
            <tbody>{kp_rows}</tbody>
        </table>"""

    encouragement = exam_data.get("encouragement", "继续加油，你正在变得越来越好！")

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{"班级错题报告" if is_teacher else "个人学习报告"}</title>
<style>
    @page {{ size: A4; margin: 20mm; }}
    * {{ box-sizing: border-box; }}
    body {{
        font-family: "Microsoft YaHei", "SimSun", sans-serif;
        color: #333;
        max-width: 210mm;
        margin: 0 auto;
        padding: 20px;
    }}
    h1 {{ text-align: center; font-size: 20px; margin-bottom: 5px; }}
    .subtitle {{ text-align: center; color: #666; font-size: 13px; margin-bottom: 20px; }}
    .summary {{ display: flex; gap: 16px; justify-content: center; margin: 16px 0; }}
    .card {{
        background: #f5f5f5; border-radius: 8px; padding: 12px 20px;
        text-align: center; min-width: 100px;
    }}
    .card .value {{ font-size: 24px; font-weight: bold; color: #e74c3c; }}
    .card .label {{ font-size: 12px; color: #888; }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0; }}
    th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; font-size: 13px; }}
    th {{ background: #f0f0f0; font-weight: bold; }}
    .encouragement {{
        margin-top: 20px; padding: 16px; background: #e8f5e9;
        border-left: 4px solid #4caf50; border-radius: 4px; font-size: 14px;
    }}
    @media print {{
        body {{ padding: 0; }}
        .no-print {{ display: none; }}
    }}
</style>
</head>
<body>
    <h1>{"班级错题分析报告" if is_teacher else "个人学习报告"}</h1>
    <p class="subtitle">
        {exam_data.get("exam_name", "考试")} |
        应到 {exam_data.get("total_students", 0)} 人 |
        实到 {exam_data.get("present_students", 0)} 人 |
        平均分 {exam_data.get("avg_score", "N/A")}
    </p>

    <div class="summary">
        <div class="card"><div class="value">{exam_data.get("avg_score", "N/A")}</div><div class="label">平均分</div></div>
        <div class="card"><div class="value">{exam_data.get("present_students", 0)}</div><div class="label">参考人数</div></div>
        <div class="card"><div class="value">{len(exam_data.get("question_stats", []))}</div><div class="label">题目数</div></div>
    </div>

    <h2>TOP5 高频错题</h2>
    <table>
        <thead><tr><th>题号</th><th>知识点</th><th>错误人数</th><th>错误率</th></tr></thead>
        <tbody>{rows_html}</tbody>
    </table>

    {teacher_section}

    <div class="encouragement">
        <strong>{"教学建议" if is_teacher else "老师寄语"}:</strong> {encouragement}
    </div>

    <button class="no-print" onclick="window.print()" style="
        position: fixed; top: 16px; right: 16px;
        padding: 10px 20px; background: #2196f3; color: white;
        border: none; border-radius: 6px; cursor: pointer; font-size: 14px;
    ">打印 / 导出 PDF</button>

    <p class="no-print" style="text-align:center;color:#999;margin-top:24px;font-size:12px;">
        生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")} | ChemAI 智辅化学
    </p>
</body>
</html>"""
